#!/usr/bin/env python3
"""Generate RPD document for LiveKit Agents project."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ============================================================
# Page setup
# ============================================================
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ============================================================
# Style definitions
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.space_before = Pt(0)
pf.line_spacing = 1.15

# Heading styles
for level, (size, color) in enumerate([
    (Pt(26), RGBColor(0x1A, 0x73, 0xE8)),  # Heading 1
    (Pt(18), RGBColor(0x1A, 0x73, 0xE8)),  # Heading 2
    (Pt(14), RGBColor(0x19, 0x67, 0xD2)),  # Heading 3
    (Pt(12), RGBColor(0x19, 0x67, 0xD2)),  # Heading 4
], start=1):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.size = size
    h.font.color.rgb = color
    h.font.bold = True
    h.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
    h.paragraph_format.space_after = Pt(6)

# ============================================================
# Helper functions
# ============================================================
def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a styled table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1A73E8')
    
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F8F9FA')
    
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    
    return table

def add_bullet(doc, text, level=0, bold_prefix=None):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(text)
        run.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p

def add_feature_card(doc, title, description, features):
    """Add a feature card section."""
    doc.add_heading(title, level=3)
    p = doc.add_paragraph(description)
    p.runs[0].font.size = Pt(10)
    for feat in features:
        add_bullet(doc, feat)

# ============================================================
# COVER PAGE
# ============================================================
# Add blank paragraphs for vertical centering
for _ in range(6):
    doc.add_paragraph()

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('LiveKit Agents')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Requirements & Positioning Document')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x5F, 0x63, 0x68)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Realtime Voice AI Agent Framework')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x80, 0x86, 0x8B)
run.italic = True

for _ in range(4):
    doc.add_paragraph()

# Metadata table
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_table = doc.add_table(rows=4, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_table.style = 'Table Grid'

meta_data = [
    ('Document Version', '1.0.0'),
    ('Date', datetime.date.today().strftime('%B %d, %Y')),
    ('License', 'Apache 2.0 (Open Source)'),
    ('Python Version', '3.10 - 3.14'),
]
for i, (key, val) in enumerate(meta_data):
    c0 = meta_table.rows[i].cells[0]
    c1 = meta_table.rows[i].cells[1]
    c0.text = ''
    c1.text = ''
    r0 = c0.paragraphs[0].add_run(key)
    r0.bold = True
    r0.font.size = Pt(11)
    r1 = c1.paragraphs[0].add_run(val)
    r1.font.size = Pt(11)
    set_cell_shading(c0, 'E8F0FE')
    c0.width = Cm(5)
    c1.width = Cm(8)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (placeholder)
# ============================================================
doc.add_heading('Table of Contents', level=1)
p = doc.add_paragraph()
run = p.add_run('[Insert TOC here - update fields in Word]')
run.italic = True
run.font.color.rgb = RGBColor(0x80, 0x86, 0x8B)

toc_items = [
    '1. Executive Summary',
    '2. Project Positioning',
    '3. Architecture Overview',
    '4. Core Capabilities',
    '5. Plugin Ecosystem & Provider Support',
    '6. Agent Types & Patterns',
    '7. Deployment Modes',
    '8. Testing & Quality Assurance',
    '9. Key Technical Innovations',
    '10. Use Cases & Applications',
    '11. Competitive Advantages',
    '12. Development & Operations',
    'Appendix A: Directory Structure',
    'Appendix B: Environment Configuration',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
doc.add_heading('1. Executive Summary', level=1)

doc.add_heading('1.1 Project Overview', level=2)
doc.add_paragraph(
    'LiveKit Agents is an open-source Python framework for building realtime, '
    'programmable voice AI participants that run on servers. It enables developers '
    'to create conversational, multi-modal voice agents that can see, hear, and '
    'understand, connected via WebRTC through the LiveKit infrastructure.'
)

doc.add_heading('1.2 Vision Statement', level=2)
doc.add_paragraph(
    'To provide the most flexible, production-ready foundation for building '
    'voice AI agents that feel natural, respond instantly, and integrate seamlessly '
    'with any AI model or external service.'
)

doc.add_heading('1.3 Key Metrics', level=2)
add_styled_table(doc,
    ['Metric', 'Value'],
    [
        ('License', 'Apache 2.0'),
        ('Language', 'Python 3.10-3.14'),
        ('Plugin Ecosystem', '60+ provider plugins'),
        ('Supported STT Providers', '12+'),
        ('Supported TTS Providers', '18+'),
        ('Supported LLM Providers', '8+'),
        ('Avatar Providers', '15+'),
        ('Deployment Modes', '4 (Console, Dev, Production, Connect)'),
    ],
    col_widths=[6, 8]
)

doc.add_page_break()

# ============================================================
# 2. PROJECT POSITIONING
# ============================================================
doc.add_heading('2. Project Positioning', level=1)

doc.add_heading('2.1 What Is LiveKit Agents?', level=2)
doc.add_paragraph(
    'LiveKit Agents bridges the gap between AI models and realtime communication. '
    'Rather than being just another chatbot framework, it is purpose-built for '
    'low-latency, bidirectional voice interactions over WebRTC, with first-class '
    'support for video, text, and tool use.'
)

doc.add_heading('2.2 Target Audience', level=2)
add_styled_table(doc,
    ['Audience', 'Use Case'],
    [
        ('AI/ML Engineers', 'Build and deploy voice agents with custom LLM pipelines'),
        ('Telecom Developers', 'Create SIP-based phone agents with IVR navigation'),
        ('Customer Service Teams', 'Automate call centers with intelligent voice bots'),
        ('Product Teams', 'Add voice interfaces to existing applications'),
        ('Researchers', 'Experiment with multimodal AI and realtime models'),
        ('System Integrators', 'Connect AI agents to existing infrastructure via MCP'),
    ],
    col_widths=[5, 9]
)

doc.add_heading('2.3 Core Value Proposition', level=2)
add_feature_card(doc,
    'Infrastructure-Agnostic',
    'Works with any LiveKit server deployment — cloud, self-hosted, or hybrid.',
    ['No vendor lock-in for infrastructure', 'Scales from single instance to distributed clusters', 'Built-in load balancing and health checks']
)
add_feature_card(doc,
    'Model-Agnostic',
    'Swap between 60+ AI providers without changing agent code.',
    ['Unified inference abstraction layer', 'Automatic fallback between providers', 'LiveKit Cloud inference routing eliminates individual API keys']
)
add_feature_card(doc,
    'Production-Ready',
    'Built for reliability at scale from day one.',
    ['OpenTelemetry observability', 'Session recording and replay', 'Graceful draining and automatic retry', 'Prometheus metrics integration']
)

doc.add_page_break()

# ============================================================
# 3. ARCHITECTURE OVERVIEW
# ============================================================
doc.add_heading('3. Architecture Overview', level=1)

doc.add_heading('3.1 Layered Architecture', level=2)
doc.add_paragraph(
    'The framework follows a layered, event-driven design with clear separation of concerns:'
)

add_styled_table(doc,
    ['Layer', 'Component', 'Responsibility'],
    [
        ('Server/Worker', 'AgentServer (Worker)', 'Orchestrates job scheduling, process pooling, load balancing, and LiveKit server communication'),
        ('Job Execution', 'JobContext / JobProcess', 'Provides execution context, manages OS process/thread lifecycle, warm pool management'),
        ('Voice Agent', 'Agent / AgentSession', 'Core runtime managing STT/LLM/TTS pipeline, turn detection, interruptions, and tool orchestration'),
        ('Model Interface', 'STT / TTS / LLM / RealtimeModel / VAD', 'Provider-agnostic abstract base classes with streaming and fallback adapters'),
        ('Inference', 'LiveKit Inference Layer', 'Unified routing of model requests through LiveKit Cloud using provider/model identifiers'),
        ('I/O', 'RoomIO / RecorderIO / AgentInput/Output', 'Bridges agent session with LiveKit room media streams and recording'),
    ],
    col_widths=[3, 4.5, 7]
)

doc.add_heading('3.2 Data Flow', level=2)
doc.add_paragraph('The voice processing pipeline follows this flow:')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'User Audio → RoomIO → AgentSession → VAD → STT → Agent.stt_node → '
    'Agent.on_user_turn_completed → Agent.llm_node → Agent.transcription_node → '
    'Agent.tts_node → Agent.realtime_audio_output_node → RoomIO → User'
)
run.font.size = Pt(9)
run.font.name = 'Consolas'
run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

doc.add_heading('3.3 IPC Architecture', level=2)
doc.add_paragraph(
    'The Inter-Process Communication system manages a warm pool of processes and threads '
    'for fast job startup, supporting both process-based ("spawn"/"forkserver") and '
    'thread-based execution models across platforms.'
)
add_bullet(doc, 'Process Pool (ProcPool) manages worker lifecycle and job dispatching')
add_bullet(doc, 'JobProcExecutor handles process-based job execution')
add_bullet(doc, 'JobThreadExecutor handles thread-based job execution')
add_bullet(doc, 'Multiprocessing contexts ensure cross-platform compatibility')

doc.add_heading('3.4 Pipeline Node Architecture', level=2)
doc.add_paragraph(
    'Each stage of the voice pipeline is implemented as an overridable node, '
    'allowing deep customization without rewriting the entire agent:'
)
add_styled_table(doc,
    ['Node', 'Purpose', 'Customization'],
    [
        ('stt_node', 'Speech-to-text processing', 'Override transcription logic, add post-processing'),
        ('llm_node', 'Language model inference', 'Custom prompt engineering, tool selection'),
        ('tts_node', 'Text-to-speech synthesis', 'Voice selection, prosody control'),
        ('transcription_node', 'Transcript management', 'Timed transcripts, word-level sync'),
        ('realtime_audio_output_node', 'Audio emission', 'Streaming control, interruption handling'),
    ],
    col_widths=[4, 4, 6.5]
)

doc.add_page_break()

# ============================================================
# 4. CORE CAPABILITIES
# ============================================================
doc.add_heading('4. Core Capabilities', level=1)

doc.add_heading('4.1 Voice Interaction', level=2)
add_feature_card(doc,
    'Semantic Turn Detection',
    'Uses transformer models to detect when a user is done speaking, reducing false interruptions.',
    ['Multilingual support', 'Adaptive sensitivity', 'VAD-based fallback']
)
add_feature_card(doc,
    'Preemptive Generation',
    'Speculatively begins LLM/TTS requests before end-of-turn is detected, dramatically reducing response latency.',
    ['Overlaps inference with user audio', 'Configurable speculation window', 'Automatic cancellation on new input']
)
add_feature_card(doc,
    'Interruption Handling',
    'Adaptive and VAD-based interruption detection with Acoustic Echo Cancellation (AEC) warmup protection.',
    ['Distinguishes meaningful interruptions from noise', 'AEC warmup prevents audio artifacts', 'Configurable interruption thresholds']
)
add_feature_card(doc,
    'Background Audio',
    'Ambient and thinking audio playback for improved realism during processing delays.',
    ['Thinking sounds during LLM inference', 'Ambient audio for idle periods', 'Smooth transitions']
)

doc.add_heading('4.2 Multi-Agent System', level=2)
add_feature_card(doc,
    'Agent Handoff',
    'Agents can transfer control to other agents, preserving chat context and enabling specialized workflows.',
    ['Seamless context preservation', 'Specialized agent routing', 'Automatic state transfer']
)
add_feature_card(doc,
    'AgentTask (Inline Sub-Agents)',
    'Sub-agents that can be awaited inside function tools for complex multi-step workflows.',
    ['Automatic context preservation', 'Chat history merging', 'Seamless handoff back to parent']
)

doc.add_heading('4.3 Tool Use & Integration', level=2)
add_feature_card(doc,
    'MCP Support',
    'Native Model Context Protocol integration for connecting to external tool servers.',
    ['stdio transport', 'SSE transport', 'Streamable HTTP transport']
)
add_feature_card(doc,
    'Dynamic Tool Creation',
    'Tools can be created and modified at runtime for adaptive behavior.',
    ['Runtime tool registration', 'Dynamic parameter schemas', 'Context-aware tool selection']
)
add_feature_card(doc,
    'Structured Output',
    'LLM structured output support for controlled, parseable responses.',
    ['JSON schema validation', 'Type-safe outputs', 'Retry on validation failure']
)

doc.add_heading('4.4 Telephony', level=2)
add_feature_card(doc,
    'SIP Integration',
    'Full SIP support for making and receiving phone calls.',
    ['Inbound call handling', 'Outbound dialing', 'Warm and cold transfers']
)
add_feature_card(doc,
    'IVR Detection',
    'Automatic detection and navigation of IVR phone tree systems.',
    ['Menu option recognition', 'Automated navigation', 'Configurable IVR rules']
)
add_feature_card(doc,
    'Answering Machine Detection',
    'Detects if a call was answered by a machine vs. human.',
    ['AMD classification', 'Configurable detection policies', 'Automated hang-up or message leaving']
)

doc.add_heading('4.5 Multimodal Support', level=2)
add_feature_card(doc,
    'Video Avatar Integration',
    'Support for synchronized video avatar providers.',
    ['15+ avatar providers supported', 'Lip-sync synchronization', 'Real-time video streaming']
)
add_feature_card(doc,
    'Video Input',
    'Voice activity-based video sampling for multimodal agents.',
    ['Automatic frame capture', 'Configurable sampling rate', 'Integration with vision LLMs']
)

doc.add_heading('4.6 Observability & Recording', level=2)
add_feature_card(doc,
    'Session Recording',
    'Granular control over recording session audio, traces, logs, and transcripts.',
    ['Audio recording to OGG format', 'Automatic upload to LiveKit Cloud', 'Configurable recording scope']
)
add_feature_card(doc,
    'OpenTelemetry Integration',
    'Full observability with trace spans for every LLM/STT/TTS request.',
    ['Automatic trace span creation', 'Langfuse compatibility', 'Prometheus metrics']
)

doc.add_page_break()

# ============================================================
# 5. PLUGIN ECOSYSTEM & PROVIDER SUPPORT
# ============================================================
doc.add_heading('5. Plugin Ecosystem & Provider Support', level=1)

doc.add_heading('5.1 Plugin Architecture', level=2)
doc.add_paragraph(
    'Plugins are managed as a uv workspace with 60+ plugins. Each plugin is a separate '
    'package following the pattern livekit-plugins-<provider>. Plugins register themselves '
    'via the Plugin base class and can provide optional download_files() for model/resource downloads.'
)
add_bullet(doc, 'Auto-registration via Plugin.register_plugin()')
add_bullet(doc, 'Optional dependency groups for selective installation')
add_bullet(doc, 'Connection pooling for efficient streaming inference')
add_bullet(doc, 'Stream adapters for non-streaming providers')
add_bullet(doc, 'Fallback adapters for resilience')

doc.add_heading('5.2 Speech-to-Text (STT) Providers', level=2)
add_styled_table(doc,
    ['Provider', 'Key Features'],
    [
        ('Deepgram', 'Nova-3 model, streaming, diarization'),
        ('Google Cloud Speech', 'Multi-language, real-time streaming'),
        ('Azure Speech', 'Neural models, custom models'),
        ('AWS Transcribe', 'Medical, call analytics'),
        ('AssemblyAI', 'Speaker labels, content moderation'),
        ('Speechmatics', 'Multi-language, diarization'),
        ('Soniox', 'Low-latency, custom vocabulary'),
        ('Gladia', 'Multi-language, code-switching'),
        ('Clova (Naver)', 'Korean language optimization'),
        ('Spitch', 'Romanian language specialist'),
        ('Simplismart', 'Enterprise-grade accuracy'),
        ('Phonic', 'Realtime STT with low latency'),
    ],
    col_widths=[5, 9]
)

doc.add_heading('5.3 Text-to-Speech (TTS) Providers', level=2)
add_styled_table(doc,
    ['Provider', 'Key Features'],
    [
        ('Cartesia', 'Sonic-3 model, ultra-low latency'),
        ('ElevenLabs', 'Voice cloning, emotional range'),
        ('OpenAI', 'TTS-1, HD models'),
        ('Google Cloud TTS', 'Neural voices, multi-language'),
        ('Azure TTS', 'Neural voices, custom voices'),
        ('LMNT', 'Expressive speech, low latency'),
        ('Rime', 'Custom voice creation'),
        ('Fish Audio', 'Open-source models'),
        ('Murf', 'Studio-quality voices'),
        ('Neuphonic', 'Realtime streaming'),
        ('Resemble', 'Voice cloning, emotion control'),
        ('Speechify', 'Natural-sounding voices'),
        ('Sarvam', 'Indian languages'),
        ('Minimax', 'Chinese/English bilingual'),
        ('Hamming', 'High-fidelity synthesis'),
        ('Baseten', 'Self-hosted deployment'),
        ('UpLift AI', 'Conversational voices'),
        ('Gradium', 'Expressive speech'),
        ('CambAI', '160+ languages'),
    ],
    col_widths=[4, 10]
)

doc.add_heading('5.4 Language Model (LLM) Providers', level=2)
add_styled_table(doc,
    ['Provider', 'Models'],
    [
        ('OpenAI', 'GPT-4o, GPT-5 series, o1, o3, o4'),
        ('Anthropic', 'Claude 3/4 series'),
        ('Google', 'Gemini 1.5/2.0'),
        ('Groq', 'Llama, Mixtral (ultra-fast inference)'),
        ('Fireworks AI', 'Open-source model hosting'),
        ('Mistral AI', 'Mistral, Mixtral'),
        ('NVIDIA', 'Nemotron, Llama on DGX Cloud'),
        ('xAI', 'Grok series'),
        ('Minimax', 'Chinese/English bilingual models'),
    ],
    col_widths=[4, 10]
)

doc.add_heading('5.5 Realtime Models', level=2)
add_styled_table(doc,
    ['Provider', 'Description'],
    [
        ('OpenAI Realtime API', 'Unified audio+text model with function calling'),
        ('Google Gemini Live', 'Multimodal realtime conversation'),
        ('Ultravox', 'Open-source realtime voice model'),
        ('Phonic Realtime', 'Low-latency voice interaction'),
        ('Personaplex', 'Personalized voice AI'),
    ],
    col_widths=[5, 9]
)

doc.add_heading('5.6 Voice Activity Detection (VAD)', level=2)
add_styled_table(doc,
    ['Provider', 'Description'],
    [
        ('Silero VAD', 'Local ONNX model, lightweight'),
        ('Turn Detector', 'Transformer-based multilingual model'),
    ],
    col_widths=[5, 9]
)

doc.add_heading('5.7 Avatar Providers', level=2)
add_styled_table(doc,
    ['Providers'],
    [
        ['Tavus, Hedra, BitHuman, LemonSlice, Simli, Anam, Avatario, Avatartalk, Bey, Keyframe, LiveAvatar, Inworld, Hume, AsyncAI'],
    ],
    col_widths=[14]
)

doc.add_page_break()

# ============================================================
# 6. AGENT TYPES & PATTERNS
# ============================================================
doc.add_heading('6. Agent Types & Patterns', level=1)

add_styled_table(doc,
    ['Agent Type', 'Description', 'Use Case'],
    [
        ('Basic Voice Agent', 'Single agent with STT/LLM/TTS pipeline', 'Simple conversational bots, FAQ agents'),
        ('Multi-Agent Handoff', 'Chain of specialized agents', 'Complex workflows requiring different expertise'),
        ('Realtime Agent', 'Using unified audio+text models', 'Ultra-low latency conversations'),
        ('Text-Only Agent', 'Skip voice for text interactions', 'Chat bots, messaging integrations'),
        ('Push-to-Talk Agent', 'Multi-user response via PTT', 'Meeting assistants, group moderation'),
        ('Avatar Agent', 'Voice agent with video avatar', 'Virtual receptionists, training bots'),
        ('Telephony Agent', 'SIP-based phone call handling', 'Call centers, appointment scheduling'),
        ('Task-Based Agent', 'Inline sub-agents in function tools', 'Multi-step workflows, research agents'),
        ('Browser Agent', 'Web automation + voice', 'Web research assistants, form filling'),
        ('RAG Agent', 'Retrieval-augmented generation', 'Knowledge base Q&A, document search'),
    ],
    col_widths=[3.5, 5, 5.5]
)

doc.add_page_break()

# ============================================================
# 7. DEPLOYMENT MODES
# ============================================================
doc.add_heading('7. Deployment Modes', level=1)

add_styled_table(doc,
    ['Mode', 'Command', 'Description', 'Best For'],
    [
        ('Console', 'python agent.py console', 'Terminal mode with local audio I/O. No server needed.', 'Quick testing, local development'),
        ('Development', 'python agent.py dev', 'Hot reload enabled, connects to LiveKit server.', 'Active development, multi-agent testing'),
        ('Production', 'python agent.py start', 'Process pooling, load balancing, health checks.', 'Production deployment, scaling'),
        ('Connect', 'python agent.py connect', 'Connect to an existing LiveKit room.', 'Debugging, ad-hoc room joining'),
    ],
    col_widths=[2.5, 4.5, 5.5, 4]
)

doc.add_heading('7.1 Production Features', level=2)
add_bullet(doc, 'Built-in HTTP health check server')
add_bullet(doc, 'Prometheus metrics endpoint')
add_bullet(doc, 'Automatic retry on connection failure')
add_bullet(doc, 'Graceful draining for zero-downtime deployments')
add_bullet(doc, 'Connection resilience with automatic reconnection')
add_bullet(doc, 'Unregistered mode for isolated/local testing')

doc.add_page_break()

# ============================================================
# 8. TESTING & QUALITY ASSURANCE
# ============================================================
doc.add_heading('8. Testing & Quality Assurance', level=1)

doc.add_heading('8.1 Testing Framework', level=2)
doc.add_paragraph(
    'The framework includes a comprehensive testing suite with multiple layers of validation:'
)

add_feature_card(doc,
    'Unit Tests',
    'Located in tests/, covering tools, audio, IPC, agent session, chat context, tokenization, and transcription filters.',
    ['Fake implementations for all model types', 'Isolated component testing', 'Fast execution']
)
add_feature_card(doc,
    'RunResult Testing',
    'The session.run() API returns a RunResult with an expect assertion interface.',
    ['Event assertions (is_message, is_function_call)', 'LLM-based judging for behavior validation', 'Event range assertions']
)
add_feature_card(doc,
    'Network Fault Injection',
    'Toxiproxy integration for testing resilience under network failures.',
    ['Simulated latency', 'Connection drops', 'Bandwidth throttling']
)
add_feature_card(doc,
    'Task Leak Detection',
    'Automatic detection of leaked asyncio tasks in tests.',
    ['Prevents resource leaks', 'Ensures clean shutdown', 'Catches async bugs early']
)
add_feature_card(doc,
    'Integration Tests',
    'Plugin-specific tests requiring real API credentials, run in GitHub CI.',
    ['End-to-end validation', 'Provider compatibility', 'CI/CD pipeline integration']
)

doc.add_heading('8.2 Code Quality', level=2)
add_styled_table(doc,
    ['Tool', 'Purpose'],
    [
        ('ruff (format)', 'Code formatting (100 char line length)'),
        ('ruff (lint)', 'Static analysis and linting'),
        ('mypy', 'Strict type checking'),
        ('pytest', 'Test execution'),
        ('make check', 'Run all checks (format, lint, type)'),
    ],
    col_widths=[5, 9]
)

doc.add_page_break()

# ============================================================
# 9. KEY TECHNICAL INNOVATIONS
# ============================================================
doc.add_heading('9. Key Technical Innovations', level=1)

doc.add_heading('9.1 Pipeline Node Architecture', level=2)
doc.add_paragraph(
    'The Agent\'s processing pipeline (stt_node, llm_node, tts_node, etc.) is fully '
    'overridable, allowing deep customization of each processing stage without rewriting '
    'the entire agent. This enables developers to inject custom logic at any point in '
    'the voice processing chain.'
)

doc.add_heading('9.2 Preemptive Generation', level=2)
doc.add_paragraph(
    'The session speculatively starts LLM and TTS requests as soon as a user transcript '
    'is received, overlapping inference with user audio to dramatically reduce response '
    'latency. This is one of the most impactful optimizations for perceived responsiveness.'
)

doc.add_heading('9.3 Adaptive Interruption Detection', level=2)
doc.add_paragraph(
    'Uses a transformer model to distinguish between meaningful user interruptions and '
    'background noise, reducing false interruptions that break the conversational flow.'
)

doc.add_heading('9.4 Inference Abstraction Layer', level=2)
doc.add_paragraph(
    'The inference.STT/LLM/TTS classes provide a unified interface to access any model '
    'provider through LiveKit Cloud using simple string identifiers (e.g., '
    '"deepgram/nova-3", "openai/gpt-4.1-mini", "cartesia/sonic-3"), abstracting away '
    'provider-specific authentication and APIs.'
)

doc.add_heading('9.5 AgentTask System', level=2)
doc.add_paragraph(
    'Inline sub-agents that can be awaited inside function tools, with automatic context '
    'preservation, chat history merging, and seamless handoff back to the parent agent. '
    'This enables complex multi-step workflows within a single conversation.'
)

doc.add_heading('9.6 Connection Pool Architecture', level=2)
doc.add_paragraph(
    'TTS plugins use pooled WebSocket connections via ConnectionPool for efficient '
    'streaming inference, reducing connection overhead and improving throughput.'
)

doc.add_heading('9.7 Stream Adapter Pattern', level=2)
doc.add_paragraph(
    'Automatic wrapping of non-streaming STT/TTS providers with streaming adapters using '
    'VAD or sentence tokenization, enabling any provider to work in the streaming pipeline.'
)

doc.add_heading('9.8 Timed Transcripts', level=2)
doc.add_paragraph(
    'Audio frames carry timestamped transcript metadata, enabling precise word-level '
    'synchronization between text and audio playback through the TranscriptSynchronizer.'
)

doc.add_page_break()

# ============================================================
# 10. USE CASES & APPLICATIONS
# ============================================================
doc.add_heading('10. Use Cases & Applications', level=1)

add_styled_table(doc,
    ['Domain', 'Application', 'Key Features Used'],
    [
        ('Customer Service', 'Intelligent call center agents', 'SIP, IVR, AMD, Multi-agent handoff, MCP tools'),
        ('Sales', 'Outbound lead qualification', 'Telephony, AMD, Structured output, Recording'),
        ('Healthcare', 'Appointment scheduling & reminders', 'SIP, IVR, Calendar integration via MCP'),
        ('Education', 'Language tutoring agents', 'Realtime models, Avatar, Preemptive generation'),
        ('Entertainment', 'Interactive voice characters', 'Avatar, Background audio, Multimodal'),
        ('Research', 'Voice AI experimentation', '60+ providers, Realtime models, Testing framework'),
        ('Enterprise', 'Internal knowledge assistants', 'RAG, MCP, Structured output, Recording'),
        ('Accessibility', 'Voice interfaces for applications', 'Text-only mode, Push-to-talk, Multi-modal'),
        ('IoT', 'Voice-controlled smart devices', 'Console mode, Local audio I/O, Low latency'),
        ('Media', 'Automated content creation', 'Browser agent, LLM, TTS with voice cloning'),
    ],
    col_widths=[3, 5, 6]
)

doc.add_page_break()

# ============================================================
# 11. COMPETITIVE ADVANTAGES
# ============================================================
doc.add_heading('11. Competitive Advantages', level=1)

add_styled_table(doc,
    ['Advantage', 'Description'],
    [
        ('Open Source', 'Apache 2.0 license — no vendor lock-in, full transparency'),
        ('Model Agnostic', '60+ providers — swap models without code changes'),
        ('Production Ready', 'Built-in observability, recording, health checks, metrics'),
        ('Low Latency', 'Preemptive generation, streaming adapters, connection pooling'),
        ('Flexible Architecture', 'Overridable pipeline nodes, custom agents, MCP integration'),
        ('Multi-Modal', 'Voice, video, text, avatars — all in one framework'),
        ('Telephony Native', 'SIP, IVR, AMD built-in — not an afterthought'),
        ('Testing Framework', 'RunResult assertions, LLM judging, fault injection'),
        ('Scalable', 'Process pooling, load balancing, distributed job execution'),
        ('Community', 'Active development, extensive examples, comprehensive documentation'),
    ],
    col_widths=[4, 10]
)

doc.add_page_break()

# ============================================================
# 12. DEVELOPMENT & OPERATIONS
# ============================================================
doc.add_heading('12. Development & Operations', level=1)

doc.add_heading('12.1 Build & Development Commands', level=2)
add_styled_table(doc,
    ['Command', 'Description'],
    [
        ('make install', 'Install all dependencies with dev extras'),
        ('make format', 'Format code with ruff'),
        ('make lint', 'Run ruff linter'),
        ('make lint-fix', 'Run ruff linter and auto-fix issues'),
        ('make type-check', 'Run mypy type checker (strict mode)'),
        ('make check', 'Run all checks (format, lint, type)'),
        ('uv run pytest', 'Run all tests'),
        ('make doctor', 'Check development environment health'),
    ],
    col_widths=[5, 9]
)

doc.add_heading('12.2 Environment Variables', level=2)
add_styled_table(doc,
    ['Variable', 'Purpose'],
    [
        ('LIVEKIT_URL', 'WebSocket URL of LiveKit server'),
        ('LIVEKIT_API_KEY', 'API key for authentication'),
        ('LIVEKIT_API_SECRET', 'API secret for authentication'),
        ('OPENAI_API_KEY', 'OpenAI API key'),
        ('DEEPGRAM_API_KEY', 'Deepgram API key'),
        ('ANTHROPIC_API_KEY', 'Anthropic API key'),
    ],
    col_widths=[6, 8]
)

doc.add_heading('12.3 Code Style Standards', level=2)
add_bullet(doc, 'Line length: 100 characters')
add_bullet(doc, 'Python 3.10+ compatibility')
add_bullet(doc, 'Google-style docstrings')
add_bullet(doc, 'Strict mypy type checking enabled')
add_bullet(doc, 'Use make check and make fix before committing')

doc.add_page_break()

# ============================================================
# APPENDIX A: DIRECTORY STRUCTURE
# ============================================================
doc.add_heading('Appendix A: Directory Structure', level=1)

p = doc.add_paragraph()
run = p.add_run(
    'livekit-agents/livekit/agents/\n'
    '├── voice/              # Core voice agent: AgentSession, Agent, room I/O, transcription\n'
    '├── llm/                # LLM integration: chat context, tool definitions, MCP support\n'
    '├── stt/                # Speech-to-text with fallback and stream adapters\n'
    '├── tts/                # Text-to-speech with fallback and stream pacing\n'
    '├── ipc/                # Inter-process communication for distributed job execution\n'
    '├── cli/                # CLI commands (console, dev, start, connect)\n'
    '├── inference/          # Remote model inference (LLM, STT, TTS)\n'
    '├── telemetry/          # OpenTelemetry traces and Prometheus metrics\n'
    '└── utils/              # Audio processing, codecs, HTTP, async utilities\n'
    '\n'
    'livekit-plugins/        # 60+ provider plugins (openai, anthropic, google, etc.)\n'
    'tests/                  # Test suite with mock implementations\n'
    'examples/               # Example agents and use cases'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_page_break()

# ============================================================
# APPENDIX B: ENVIRONMENT CONFIGURATION
# ============================================================
doc.add_heading('Appendix B: Environment Configuration', level=1)

doc.add_heading('B.1 Minimal Setup', level=2)
doc.add_paragraph(
    'For basic voice agent development, you need:'
)
add_bullet(doc, 'Python 3.10-3.14')
add_bullet(doc, 'uv package manager')
add_bullet(doc, 'LIVEKIT_URL, LIVEKIT_API_KEY, LIVEKIT_API_SECRET')
add_bullet(doc, 'At least one provider API key (e.g., OPENAI_API_KEY)')

doc.add_heading('B.2 Recommended Setup', level=2)
add_bullet(doc, 'LiveKit server (cloud or self-hosted)')
add_bullet(doc, 'Multiple STT/TTS/LLM providers for fallback')
add_bullet(doc, 'Silero VAD or Turn Detector for voice activity detection')
add_bullet(doc, 'Prometheus + Grafana for metrics visualization')
add_bullet(doc, 'Langfuse or similar for trace visualization')

doc.add_heading('B.3 Plugin Installation', level=2)
p = doc.add_paragraph()
run = p.add_run('pip install "livekit-agents[openai,silero,deepgram,cartesia,turn-detector]"')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph(
    'Optional dependencies are installed via groups, allowing you to include only the '
    'providers you need. Each plugin follows the pattern livekit-plugins-<provider>.'
)

# ============================================================
# Save document
# ============================================================
output_path = r'D:\Github\callagents\LiveKit_Agents_RPD.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
