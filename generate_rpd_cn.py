#!/usr/bin/env python3
"""生成 LiveKit Agents 项目的 RPD 文档（中文版）"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ============================================================
# 页面设置
# ============================================================
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ============================================================
# 样式定义
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
# 设置中文字体
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.space_before = Pt(0)
pf.line_spacing = 1.15

# 标题样式
for level, (size, color) in enumerate([
    (Pt(26), RGBColor(0x1A, 0x73, 0xE8)),
    (Pt(18), RGBColor(0x1A, 0x73, 0xE8)),
    (Pt(14), RGBColor(0x19, 0x67, 0xD2)),
    (Pt(12), RGBColor(0x19, 0x67, 0xD2)),
], start=1):
    h = doc.styles[f'Heading {level}']
    h.font.name = '微软雅黑'
    h.font.size = size
    h.font.color.rgb = color
    h.font.bold = True
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
    h.paragraph_format.space_after = Pt(6)

# ============================================================
# 辅助函数
# ============================================================
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1A73E8')
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F8F9FA')
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    return table

def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(text)
        run.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p

def add_feature_card(doc, title, description, features):
    doc.add_heading(title, level=3)
    p = doc.add_paragraph(description)
    p.runs[0].font.size = Pt(10)
    for feat in features:
        add_bullet(doc, feat)

# ============================================================
# 封面
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('LiveKit Agents')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('需求与定位文档')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x5F, 0x63, 0x68)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('实时语音 AI Agent 框架')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x80, 0x86, 0x8B)
run.italic = True

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_table = doc.add_table(rows=5, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_table.style = 'Table Grid'
meta_data = [
    ('文档版本', '1.0.0'),
    ('日期', datetime.date.today().strftime('%Y年%m月%d日')),
    ('许可证', 'Apache 2.0（开源）'),
    ('Python 版本', '3.10 - 3.14'),
    ('文档语言', '中文'),
]
for i, (key, val) in enumerate(meta_data):
    c0 = meta_table.rows[i].cells[0]
    c1 = meta_table.rows[i].cells[1]
    c0.text = ''
    c1.text = ''
    r0 = c0.paragraphs[0].add_run(key)
    r0.bold = True
    r0.font.size = Pt(11)
    r1 = c1.paragraphs[0].add_run(val)
    r1.font.size = Pt(11)
    set_cell_shading(c0, 'E8F0FE')
    c0.width = Cm(5)
    c1.width = Cm(8)

doc.add_page_break()

# ============================================================
# 目录
# ============================================================
doc.add_heading('目录', level=1)
p = doc.add_paragraph()
run = p.add_run('[请在 Word 中更新域以生成自动目录]')
run.italic = True
run.font.color.rgb = RGBColor(0x80, 0x86, 0x8B)

toc_items = [
    '1. 执行摘要',
    '2. 项目定位',
    '3. 架构概览',
    '4. 核心能力',
    '5. 插件生态与供应商支持',
    '6. Agent 类型与模式',
    '7. 部署模式',
    '8. 测试与质量保证',
    '9. 关键技术创新',
    '10. 应用场景',
    '11. 竞争优势',
    '12. 开发与运维',
    '附录 A：目录结构',
    '附录 B：环境配置',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. 执行摘要
# ============================================================
doc.add_heading('1. 执行摘要', level=1)

doc.add_heading('1.1 项目概述', level=2)
doc.add_paragraph(
    'LiveKit Agents 是一个开源的 Python 框架，用于构建运行在服务器上的实时、可编程的语音 AI 参与者。'
    '它使开发者能够创建可通过 WebRTC 与 LiveKit 基础设施连接的对话式、多模态语音 Agent，'
    '这些 Agent 能够"看"、"听"和"理解"。'
)

doc.add_heading('1.2 愿景声明', level=2)
doc.add_paragraph(
    '为构建自然交互、即时响应、且能与任何 AI 模型或外部服务无缝集成的语音 AI Agent，'
    '提供最灵活、最可用于生产的底层框架。'
)

doc.add_heading('1.3 关键指标', level=2)
add_styled_table(doc,
    ['指标', '数值'],
    [
        ('许可证', 'Apache 2.0'),
        ('语言', 'Python 3.10-3.14'),
        ('插件生态', '60+ 供应商插件'),
        ('支持的 STT 供应商', '12+'),
        ('支持的 TTS 供应商', '18+'),
        ('支持的 LLM 供应商', '8+'),
        ('数字人供应商', '15+'),
        ('部署模式', '4 种（控制台、开发、生产、连接）'),
    ],
    col_widths=[6, 8]
)

doc.add_page_break()

# ============================================================
# 2. 项目定位
# ============================================================
doc.add_heading('2. 项目定位', level=1)

doc.add_heading('2.1 什么是 LiveKit Agents？', level=2)
doc.add_paragraph(
    'LiveKit Agents 架起了 AI 模型与实时通信之间的桥梁。它不仅仅是一个聊天机器人框架，'
    '而是专为基于 WebRTC 的低延迟、双向语音交互而设计，并原生支持视频、文本和工具调用。'
)

doc.add_heading('2.2 目标用户', level=2)
add_styled_table(doc,
    ['用户群体', '应用场景'],
    [
        ('AI/ML 工程师', '构建和部署带有自定义 LLM 流水线的语音 Agent'),
        ('电信开发者', '创建基于 SIP 的电话 Agent，支持 IVR 导航'),
        ('客服团队', '用智能语音机器人自动化呼叫中心'),
        ('产品团队', '为现有应用添加语音交互界面'),
        ('研究人员', '实验多模态 AI 和实时模型'),
        ('系统集成商', '通过 MCP 将 AI Agent 连接到现有基础设施'),
    ],
    col_widths=[5, 9]
)

doc.add_heading('2.3 核心价值主张', level=2)

add_feature_card(doc,
    '基础设施无关',
    '兼容任何 LiveKit 服务器部署——云端、自托管或混合部署。',
    ['无基础设施供应商锁定', '从单实例到分布式集群均可扩展', '内置负载均衡与健康检查']
)
add_feature_card(doc,
    '模型无关',
    '可在 60+ AI 供应商之间切换，无需修改 Agent 代码。',
    ['统一的推理抽象层', '供应商间自动回退', 'LiveKit Cloud 推理路由消除独立 API 密钥管理']
)
add_feature_card(doc,
    '生产就绪',
    '从第一天起就为大规模可靠性而构建。',
    ['OpenTelemetry 可观测性', '会话录制与回放', '优雅排空与自动重试', 'Prometheus 指标集成']
)

doc.add_page_break()

# ============================================================
# 3. 架构概览
# ============================================================
doc.add_heading('3. 架构概览', level=1)

doc.add_heading('3.1 分层架构', level=2)
doc.add_paragraph('框架遵循分层、事件驱动的设计，关注点清晰分离：')

add_styled_table(doc,
    ['层级', '组件', '职责'],
    [
        ('服务器/工作器', 'AgentServer (Worker)', '协调任务调度、进程池、负载均衡和 LiveKit 服务器通信'),
        ('任务执行', 'JobContext / JobProcess', '提供执行上下文，管理 OS 进程/线程生命周期，热池管理'),
        ('语音 Agent', 'Agent / AgentSession', '核心运行时，管理 STT/LLM/TTS 流水线、轮次检测、打断和工具编排'),
        ('模型接口', 'STT / TTS / LLM / RealtimeModel / VAD', '供应商无关的抽象基类，支持流式和回退适配器'),
        ('推理层', 'LiveKit 推理层', '使用 provider/model 标识符通过 LiveKit Cloud 统一路由模型请求'),
        ('I/O 层', 'RoomIO / RecorderIO / AgentInput/Output', '桥接 Agent 会话与 LiveKit 房间媒体流及录制'),
    ],
    col_widths=[3, 4.5, 7]
)

doc.add_heading('3.2 数据流', level=2)
doc.add_paragraph('语音处理流水线遵循以下流程：')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    '用户音频 → RoomIO → AgentSession → VAD → STT → Agent.stt_node → '
    'Agent.on_user_turn_completed → Agent.llm_node → Agent.transcription_node → '
    'Agent.tts_node → Agent.realtime_audio_output_node → RoomIO → 用户'
)
run.font.size = Pt(9)
run.font.name = 'Consolas'
run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

doc.add_heading('3.3 IPC 架构', level=2)
doc.add_paragraph(
    '进程间通信系统管理进程和线程的热池以实现快速任务启动，'
    '支持跨平台的基于进程（"spawn"/"forkserver"）和基于线程的执行模型。'
)
add_bullet(doc, '进程池（ProcPool）管理工作器生命周期和任务分发')
add_bullet(doc, 'JobProcExecutor 处理基于进程的任务执行')
add_bullet(doc, 'JobThreadExecutor 处理基于线程的任务执行')
add_bullet(doc, '多进程上下文确保跨平台兼容性')

doc.add_heading('3.4 流水线节点架构', level=2)
doc.add_paragraph(
    '语音流水线的每个阶段都实现为可覆盖的节点，'
    '允许深度定制而无需重写整个 Agent：'
)
add_styled_table(doc,
    ['节点', '用途', '定制能力'],
    [
        ('stt_node', '语音转文字处理', '覆盖转录逻辑，添加后处理'),
        ('llm_node', '语言模型推理', '自定义提示工程，工具选择'),
        ('tts_node', '文字转语音合成', '声音选择，韵律控制'),
        ('transcription_node', '转录管理', '定时转录，词级同步'),
        ('realtime_audio_output_node', '音频输出', '流控制，打断处理'),
    ],
    col_widths=[4.5, 4, 6]
)

doc.add_page_break()

# ============================================================
# 4. 核心能力
# ============================================================
doc.add_heading('4. 核心能力', level=1)

doc.add_heading('4.1 语音交互', level=2)
add_feature_card(doc,
    '语义轮次检测',
    '使用 Transformer 模型检测用户何时说完，减少误打断。',
    ['多语言支持', '自适应灵敏度', 'VAD 回退机制']
)
add_feature_card(doc,
    '预生成（Preemptive Generation）',
    '在检测到轮次结束之前就开始 LLM/TTS 请求，大幅降低响应延迟。',
    ['推理与用户音频重叠执行', '可配置的预测窗口', '新输入时自动取消']
)
add_feature_card(doc,
    '打断处理',
    '自适应和基于 VAD 的打断检测，带有声学回声消除（AEC）预热保护。',
    ['区分有意义的打断与背景噪音', 'AEC 预热防止音频伪影', '可配置的打断阈值']
)
add_feature_card(doc,
    '背景音频',
    '在处理延迟期间播放环境和思考音频，提升真实感。',
    ['LLM 推理期间播放思考音', '空闲时段播放环境音', '平滑过渡']
)

doc.add_heading('4.2 多 Agent 系统', level=2)
add_feature_card(doc,
    'Agent 交接',
    'Agent 可以将控制权转移给其他 Agent，保留聊天上下文，实现专业化工作流。',
    ['无缝上下文保留', '专业化 Agent 路由', '自动状态转移']
)
add_feature_card(doc,
    'AgentTask（内联子 Agent）',
    '可在函数工具中等待的子 Agent，用于复杂的多步骤工作流。',
    ['自动上下文保留', '聊天历史合并', '无缝交还给父 Agent']
)

doc.add_heading('4.3 工具使用与集成', level=2)
add_feature_card(doc,
    'MCP 支持',
    '原生模型上下文协议集成，用于连接外部工具服务器。',
    ['stdio 传输', 'SSE 传输', '可流式 HTTP 传输']
)
add_feature_card(doc,
    '动态工具创建',
    '工具可在运行时创建和修改，实现自适应行为。',
    ['运行时工具注册', '动态参数模式', '上下文感知工具选择']
)
add_feature_card(doc,
    '结构化输出',
    'LLM 结构化输出支持，实现可控、可解析的响应。',
    ['JSON Schema 验证', '类型安全输出', '验证失败自动重试']
)

doc.add_heading('4.4 电话通信', level=2)
add_feature_card(doc,
    'SIP 集成',
    '完整的 SIP 支持，用于拨打和接听电话。',
    ['来电处理', '外拨', '热转接和冷转接']
)
add_feature_card(doc,
    'IVR 检测',
    '自动检测和导航 IVR 电话树系统。',
    ['菜单选项识别', '自动导航', '可配置的 IVR 规则']
)
add_feature_card(doc,
    '应答机检测（AMD）',
    '检测来电是由机器还是真人接听。',
    ['AMD 分类', '可配置的检测策略', '自动挂断或留言']
)

doc.add_heading('4.5 多模态支持', level=2)
add_feature_card(doc,
    '数字人集成',
    '支持同步视频数字人供应商。',
    ['支持 15+ 数字人供应商', '唇形同步', '实时视频流']
)
add_feature_card(doc,
    '视频输入',
    '基于语音活动的视频采样，用于多模态 Agent。',
    ['自动帧捕获', '可配置采样率', '与视觉 LLM 集成']
)

doc.add_heading('4.6 可观测性与录制', level=2)
add_feature_card(doc,
    '会话录制',
    '细粒度控制录制会话音频、追踪、日志和转录。',
    ['OGG 格式音频录制', '自动上传至 LiveKit Cloud', '可配置的录制范围']
)
add_feature_card(doc,
    'OpenTelemetry 集成',
    '完整的可观测性，每个 LLM/STT/TTS 请求都有追踪跨度。',
    ['自动创建追踪跨度', 'Langfuse 兼容', 'Prometheus 指标']
)

doc.add_page_break()

# ============================================================
# 5. 插件生态与供应商支持
# ============================================================
doc.add_heading('5. 插件生态与供应商支持', level=1)

doc.add_heading('5.1 插件架构', level=2)
doc.add_paragraph(
    '插件作为 uv 工作区管理，拥有 60+ 个插件。每个插件都是独立的包，'
    '遵循 livekit-plugins-<供应商> 的命名模式。插件通过 Plugin 基类自我注册，'
    '可提供可选的 download_files() 用于模型/资源下载。'
)
add_bullet(doc, '通过 Plugin.register_plugin() 自动注册')
add_bullet(doc, '可选依赖组，支持选择性安装')
add_bullet(doc, '连接池用于高效的流式推理')
add_bullet(doc, '非流式供应商的流适配器')
add_bullet(doc, '回退适配器增强韧性')

doc.add_heading('5.2 语音转文字（STT）供应商', level=2)
add_styled_table(doc,
    ['供应商', '关键特性'],
    [
        ('Deepgram', 'Nova-3 模型，流式，说话人分离'),
        ('Google Cloud Speech', '多语言，实时流式'),
        ('Azure Speech', '神经模型，自定义模型'),
        ('AWS Transcribe', '医疗，通话分析'),
        ('AssemblyAI', '说话人标签，内容审核'),
        ('Speechmatics', '多语言，说话人分离'),
        ('Soniox', '低延迟，自定义词汇'),
        ('Gladia', '多语言，代码切换'),
        ('Clova (Naver)', '韩语优化'),
        ('Spitch', '罗马尼亚语专精'),
        ('Simplismart', '企业级精度'),
        ('Phonic', '低延迟实时 STT'),
    ],
    col_widths=[5, 9]
)

doc.add_heading('5.3 文字转语音（TTS）供应商', level=2)
add_styled_table(doc,
    ['供应商', '关键特性'],
    [
        ('Cartesia', 'Sonic-3 模型，超低延迟'),
        ('ElevenLabs', '声音克隆，情感范围'),
        ('OpenAI', 'TTS-1，HD 模型'),
        ('Google Cloud TTS', '神经声音，多语言'),
        ('Azure TTS', '神经声音，自定义声音'),
        ('LMNT', '表现力语音，低延迟'),
        ('Rime', '自定义声音创建'),
        ('Fish Audio', '开源模型'),
        ('Murf', '录音棚级声音'),
        ('Neuphonic', '实时流式'),
        ('Resemble', '声音克隆，情感控制'),
        ('Speechify', '自然声音'),
        ('Sarvam', '印度语言'),
        ('Minimax', '中英双语'),
        ('Hamming', '高保真合成'),
        ('Baseten', '自托管部署'),
        ('UpLift AI', '对话式声音'),
        ('Gradium', '表现力语音'),
        ('CambAI', '160+ 语言'),
    ],
    col_widths=[4, 10]
)

doc.add_heading('5.4 语言模型（LLM）供应商', level=2)
add_styled_table(doc,
    ['供应商', '模型'],
    [
        ('OpenAI', 'GPT-4o，GPT-5 系列，o1，o3，o4'),
        ('Anthropic', 'Claude 3/4 系列'),
        ('Google', 'Gemini 1.5/2.0'),
        ('Groq', 'Llama，Mixtral（超快推理）'),
        ('Fireworks AI', '开源模型托管'),
        ('Mistral AI', 'Mistral，Mixtral'),
        ('NVIDIA', 'Nemotron，DGX Cloud 上的 Llama'),
        ('xAI', 'Grok 系列'),
        ('Minimax', '中英双语模型'),
    ],
    col_widths=[4, 10]
)

doc.add_heading('5.5 实时模型', level=2)
add_styled_table(doc,
    ['供应商', '描述'],
    [
        ('OpenAI Realtime API', '统一音频+文本模型，支持函数调用'),
        ('Google Gemini Live', '多模态实时对话'),
        ('Ultravox', '开源实时语音模型'),
        ('Phonic Realtime', '低延迟语音交互'),
        ('Personaplex', '个性化语音 AI'),
    ],
    col_widths=[5, 9]
)

doc.add_heading('5.6 语音活动检测（VAD）', level=2)
add_styled_table(doc,
    ['供应商', '描述'],
    [
        ('Silero VAD', '本地 ONNX 模型，轻量级'),
        ('Turn Detector', '基于 Transformer 的多语言模型'),
    ],
    col_widths=[5, 9]
)

doc.add_heading('5.7 数字人供应商', level=2)
add_styled_table(doc,
    ['供应商列表'],
    [
        ['Tavus, Hedra, BitHuman, LemonSlice, Simli, Anam, Avatario, Avatartalk, Bey, Keyframe, LiveAvatar, Inworld, Hume, AsyncAI'],
    ],
    col_widths=[14]
)

doc.add_page_break()

# ============================================================
# 6. Agent 类型与模式
# ============================================================
doc.add_heading('6. Agent 类型与模式', level=1)

add_styled_table(doc,
    ['Agent 类型', '描述', '应用场景'],
    [
        ('基础语音 Agent', '单个 Agent，STT/LLM/TTS 流水线', '简单对话机器人，FAQ Agent'),
        ('多 Agent 交接', '专业化 Agent 链', '需要不同专业知识的复杂工作流'),
        ('实时 Agent', '使用统一音频+文本模型', '超低延迟对话'),
        ('纯文本 Agent', '跳过语音，仅文本交互', '聊天机器人，消息集成'),
        ('按键通话 Agent', '多用户通过按键响应', '会议助手，群组管理'),
        ('数字人 Agent', '带视频数字人的语音 Agent', '虚拟接待员，培训机器人'),
        ('电话 Agent', '基于 SIP 的电话处理', '呼叫中心，预约安排'),
        ('任务型 Agent', '函数工具中的内联子 Agent', '多步骤工作流，研究 Agent'),
        ('浏览器 Agent', '网页自动化 + 语音', '网页研究助手，表单填写'),
        ('RAG Agent', '检索增强生成', '知识库问答，文档搜索'),
    ],
    col_widths=[3.5, 5, 5.5]
)

doc.add_page_break()

# ============================================================
# 7. 部署模式
# ============================================================
doc.add_heading('7. 部署模式', level=1)

add_styled_table(doc,
    ['模式', '命令', '描述', '适用场景'],
    [
        ('控制台', 'python agent.py console', '终端模式，本地音频 I/O，无需服务器', '快速测试，本地开发'),
        ('开发', 'python agent.py dev', '启用热重载，连接 LiveKit 服务器', '活跃开发，多 Agent 测试'),
        ('生产', 'python agent.py start', '进程池，负载均衡，健康检查', '生产部署，扩展'),
        ('连接', 'python agent.py connect', '连接到现有 LiveKit 房间', '调试，临时加入房间'),
    ],
    col_widths=[2, 4.5, 5.5, 4.5]
)

doc.add_heading('7.1 生产特性', level=2)
add_bullet(doc, '内置 HTTP 健康检查服务器')
add_bullet(doc, 'Prometheus 指标端点')
add_bullet(doc, '连接失败自动重试')
add_bullet(doc, '优雅排空，实现零停机部署')
add_bullet(doc, '连接弹性，自动重连')
add_bullet(doc, '未注册模式，用于隔离/本地测试')

doc.add_page_break()

# ============================================================
# 8. 测试与质量保证
# ============================================================
doc.add_heading('8. 测试与质量保证', level=1)

doc.add_heading('8.1 测试框架', level=2)
doc.add_paragraph('框架包含多层验证的综合测试套件：')

add_feature_card(doc,
    '单元测试',
    '位于 tests/ 目录，覆盖工具、音频、IPC、Agent 会话、聊天上下文、分词和转录过滤器。',
    ['所有模型类型的 Fake 实现', '隔离的组件测试', '快速执行']
)
add_feature_card(doc,
    'RunResult 测试',
    'session.run() API 返回带有 expect 断言接口的 RunResult。',
    ['事件断言（is_message, is_function_call）', '基于 LLM 的行为评判', '事件范围断言']
)
add_feature_card(doc,
    '网络故障注入',
    'Toxiproxy 集成，测试网络故障下的韧性。',
    ['模拟延迟', '连接断开', '带宽限制']
)
add_feature_card(doc,
    '任务泄漏检测',
    '自动检测测试中泄漏的 asyncio 任务。',
    ['防止资源泄漏', '确保干净关闭', '早期捕获异步 bug']
)
add_feature_card(doc,
    '集成测试',
    '需要真实 API 凭据的插件特定测试，在 GitHub CI 中运行。',
    ['端到端验证', '供应商兼容性', 'CI/CD 流水线集成']
)

doc.add_heading('8.2 代码质量', level=2)
add_styled_table(doc,
    ['工具', '用途'],
    [
        ('ruff (format)', '代码格式化（100 字符行长度）'),
        ('ruff (lint)', '静态分析和代码检查'),
        ('mypy', '严格类型检查'),
        ('pytest', '测试执行'),
        ('make check', '运行所有检查（格式化、检查、类型）'),
    ],
    col_widths=[5, 9]
)

doc.add_page_break()

# ============================================================
# 9. 关键技术创新
# ============================================================
doc.add_heading('9. 关键技术创新', level=1)

doc.add_heading('9.1 流水线节点架构', level=2)
doc.add_paragraph(
    'Agent 的处理流水线（stt_node、llm_node、tts_node 等）完全可覆盖，'
    '允许深度定制每个处理阶段而无需重写整个 Agent。这使得开发者能够在'
    '语音处理链的任何点注入自定义逻辑。'
)

doc.add_heading('9.2 预生成（Preemptive Generation）', level=2)
doc.add_paragraph(
    '一旦收到用户转录，会话就会推测性地启动 LLM 和 TTS 请求，'
    '将推理与用户音频重叠执行，大幅降低响应延迟。'
    '这是对感知响应速度影响最大的优化之一。'
)

doc.add_heading('9.3 自适应打断检测', level=2)
doc.add_paragraph(
    '使用 Transformer 模型区分有意义的用户打断与背景噪音，'
    '减少破坏对话流畅性的误打断。'
)

doc.add_heading('9.4 推理抽象层', level=2)
doc.add_paragraph(
    'inference.STT/LLM/TTS 类提供统一接口，通过简单的字符串标识符'
    '（如 "deepgram/nova-3"、"openai/gpt-4.1-mini"、"cartesia/sonic-3"）'
    '即可通过 LiveKit Cloud 访问任何模型供应商，抽象掉供应商特定的认证和 API。'
)

doc.add_heading('9.5 AgentTask 系统', level=2)
doc.add_paragraph(
    '可在函数工具中等待的内联子 Agent，具有自动上下文保留、'
    '聊天历史合并和无缝交还给父 Agent 的能力。这使得在单次对话中'
    '实现复杂的多步骤工作流成为可能。'
)

doc.add_heading('9.6 连接池架构', level=2)
doc.add_paragraph(
    'TTS 插件通过 ConnectionPool 使用池化的 WebSocket 连接进行高效的流式推理，'
    '减少连接开销并提高吞吐量。'
)

doc.add_heading('9.7 流适配器模式', level=2)
doc.add_paragraph(
    '自动使用 VAD 或句子分词将非流式 STT/TTS 供应商包装为流式适配器，'
    '使任何供应商都能在流式流水线中工作。'
)

doc.add_heading('9.8 定时转录', level=2)
doc.add_paragraph(
    '音频帧携带带时间戳的转录元数据，通过 TranscriptSynchronizer '
    '实现文本与音频播放之间的精确词级同步。'
)

doc.add_page_break()

# ============================================================
# 10. 应用场景
# ============================================================
doc.add_heading('10. 应用场景', level=1)

add_styled_table(doc,
    ['领域', '应用', '使用的关键特性'],
    [
        ('客户服务', '智能呼叫中心 Agent', 'SIP, IVR, AMD, 多 Agent 交接, MCP 工具'),
        ('销售', '外呼线索筛选', '电话, AMD, 结构化输出, 录制'),
        ('医疗', '预约安排与提醒', 'SIP, IVR, 通过 MCP 集成日历'),
        ('教育', '语言辅导 Agent', '实时模型, 数字人, 预生成'),
        ('娱乐', '互动语音角色', '数字人, 背景音频, 多模态'),
        ('研究', '语音 AI 实验', '60+ 供应商, 实时模型, 测试框架'),
        ('企业', '内部知识助手', 'RAG, MCP, 结构化输出, 录制'),
        ('无障碍', '应用的语音界面', '纯文本模式, 按键通话, 多模态'),
        ('物联网', '语音控制智能设备', '控制台模式, 本地音频 I/O, 低延迟'),
        ('媒体', '自动化内容创作', '浏览器 Agent, LLM, 声音克隆 TTS'),
    ],
    col_widths=[3, 5, 6]
)

doc.add_page_break()

# ============================================================
# 11. 竞争优势
# ============================================================
doc.add_heading('11. 竞争优势', level=1)

add_styled_table(doc,
    ['优势', '描述'],
    [
        ('开源', 'Apache 2.0 许可证——无供应商锁定，完全透明'),
        ('模型无关', '60+ 供应商——无需修改代码即可切换模型'),
        ('生产就绪', '内置可观测性、录制、健康检查、指标'),
        ('低延迟', '预生成、流适配器、连接池'),
        ('灵活架构', '可覆盖的流水线节点、自定义 Agent、MCP 集成'),
        ('多模态', '语音、视频、文本、数字人——全部在一个框架中'),
        ('电话原生', 'SIP、IVR、AMD 内置——不是事后补充'),
        ('测试框架', 'RunResult 断言、LLM 评判、故障注入'),
        ('可扩展', '进程池、负载均衡、分布式任务执行'),
        ('社区', '活跃开发、丰富的示例、全面的文档'),
    ],
    col_widths=[4, 10]
)

doc.add_page_break()

# ============================================================
# 12. 开发与运维
# ============================================================
doc.add_heading('12. 开发与运维', level=1)

doc.add_heading('12.1 构建与开发命令', level=2)
add_styled_table(doc,
    ['命令', '描述'],
    [
        ('make install', '安装所有依赖（含开发依赖）'),
        ('make format', '使用 ruff 格式化代码'),
        ('make lint', '运行 ruff 代码检查'),
        ('make lint-fix', '运行 ruff 检查并自动修复问题'),
        ('make type-check', '运行 mypy 严格类型检查'),
        ('make check', '运行所有检查（格式化、检查、类型）'),
        ('uv run pytest', '运行所有测试'),
        ('make doctor', '检查开发环境健康状态'),
    ],
    col_widths=[5, 9]
)

doc.add_heading('12.2 环境变量', level=2)
add_styled_table(doc,
    ['变量', '用途'],
    [
        ('LIVEKIT_URL', 'LiveKit 服务器的 WebSocket URL'),
        ('LIVEKIT_API_KEY', '认证用的 API 密钥'),
        ('LIVEKIT_API_SECRET', '认证用的 API 密钥'),
        ('OPENAI_API_KEY', 'OpenAI API 密钥'),
        ('DEEPGRAM_API_KEY', 'Deepgram API 密钥'),
        ('ANTHROPIC_API_KEY', 'Anthropic API 密钥'),
    ],
    col_widths=[6, 8]
)

doc.add_heading('12.3 代码风格规范', level=2)
add_bullet(doc, '行长度：100 字符')
add_bullet(doc, 'Python 3.10+ 兼容')
add_bullet(doc, 'Google 风格文档字符串')
add_bullet(doc, '启用严格 mypy 类型检查')
add_bullet(doc, '提交前使用 make check 和 make fix')

doc.add_page_break()

# ============================================================
# 附录 A：目录结构
# ============================================================
doc.add_heading('附录 A：目录结构', level=1)

p = doc.add_paragraph()
run = p.add_run(
    'livekit-agents/livekit/agents/\n'
    '├── voice/              # 核心语音 Agent：AgentSession、Agent、房间 I/O、转录\n'
    '├── llm/                # LLM 集成：聊天上下文、工具定义、MCP 支持\n'
    '├── stt/                # 语音转文字，带回退和流适配器\n'
    '├── tts/                # 文字转语音，带回退和流节奏控制\n'
    '├── ipc/                # 进程间通信，用于分布式任务执行\n'
    '├── cli/                # CLI 命令（console、dev、start、connect）\n'
    '├── inference/          # 远程模型推理（LLM、STT、TTS）\n'
    '├── telemetry/          # OpenTelemetry 追踪和 Prometheus 指标\n'
    '└── utils/              # 音频处理、编解码器、HTTP、异步工具\n'
    '\n'
    'livekit-plugins/        # 60+ 供应商插件（openai、anthropic、google 等）\n'
    'tests/                  # 测试套件，含 Mock 实现\n'
    'examples/               # 示例 Agent 和用例'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_page_break()

# ============================================================
# 附录 B：环境配置
# ============================================================
doc.add_heading('附录 B：环境配置', level=1)

doc.add_heading('B.1 最小化设置', level=2)
doc.add_paragraph('基础语音 Agent 开发需要：')
add_bullet(doc, 'Python 3.10-3.14')
add_bullet(doc, 'uv 包管理器')
add_bullet(doc, 'LIVEKIT_URL、LIVEKIT_API_KEY、LIVEKIT_API_SECRET')
add_bullet(doc, '至少一个供应商 API 密钥（如 OPENAI_API_KEY）')

doc.add_heading('B.2 推荐设置', level=2)
add_bullet(doc, 'LiveKit 服务器（云端或自托管）')
add_bullet(doc, '多个 STT/TTS/LLM 供应商用于回退')
add_bullet(doc, 'Silero VAD 或 Turn Detector 用于语音活动检测')
add_bullet(doc, 'Prometheus + Grafana 用于指标可视化')
add_bullet(doc, 'Langfuse 或类似工具用于追踪可视化')

doc.add_heading('B.3 插件安装', level=2)
p = doc.add_paragraph()
run = p.add_run('pip install "livekit-agents[openai,silero,deepgram,cartesia,turn-detector]"')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph(
    '可选依赖通过依赖组安装，允许你只包含需要的供应商。'
    '每个插件遵循 livekit-plugins-<供应商> 的命名模式。'
)

# ============================================================
# 保存文档
# ============================================================
output_path = r'D:\Github\callagents\LiveKit_Agents_RPD_中文版.docx'
doc.save(output_path)
print(f'文档已保存至: {output_path}')
